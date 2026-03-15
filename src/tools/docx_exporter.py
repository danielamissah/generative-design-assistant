"""
Word Document Exporter for Generative Design Feasibility Report.
Uses python-docx for pure Python DOCX generation.
"""

import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger


BLUE       = RGBColor(0x2E, 0x75, 0xB6)
DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY       = RGBColor(0x66, 0x66, 0x66)


def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level == 1 else BLUE
    run.font.name = "Calibri"
    return p


def add_body(doc, text: str, bold: bool = False, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(clean(text))
    run.font.size  = Pt(10.5)
    run.font.name  = "Calibri"
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text: str):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(clean(text))
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    return p


def clean(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', str(text))
    text = re.sub(r'\*(.+?)\*',    r'\1', text)
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()


def add_two_col_table(doc, data: list, header_bg="1F4E79", row_bgs=("FFFFFF", "F0F4F8")):
    """Add a two-column key-value table."""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (key, val) in enumerate(data):
        row = table.rows[i]
        # Key cell
        kc = row.cells[0]
        kc.width = Inches(2.0)
        kr = kc.paragraphs[0].add_run(clean(str(key)))
        kr.font.bold  = True
        kr.font.size  = Pt(9.5)
        kr.font.name  = "Calibri"

        # Value cell
        vc = row.cells[1]
        vc.width = Inches(4.5)
        vr = vc.paragraphs[0].add_run(clean(str(val)))
        vr.font.size = Pt(9.5)
        vr.font.name = "Calibri"

        # Colouring
        if i == 0:
            set_cell_bg(kc, header_bg)
            set_cell_bg(vc, header_bg)
            kr.font.color.rgb = WHITE
            vr.font.color.rgb = WHITE
        else:
            bg = row_bgs[(i - 1) % 2]
            set_cell_bg(kc, "D5E8F0")
            set_cell_bg(vc, bg)

    return table


def export_docx(result: dict, output_path: str = None) -> str:
    """Generate a Word document from the agent pipeline result."""
    req  = result.get("requirements", {})
    alts = result.get("alternatives", [])
    ev   = result.get("evaluation", {})
    rep  = result.get("report", {})

    if not output_path:
        ts          = datetime.now().strftime("%Y%m%d%H%M%S")
        report_id   = rep.get("report_id", f"GDA-{ts}")
        output_path = f"outputs/reports/{report_id}.docx"

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover ──────────────────────────────────────────────────────────
    cover_label = doc.add_paragraph()
    lr = cover_label.add_run("GENERATIVE DESIGN FEASIBILITY REPORT")
    lr.font.size  = Pt(9)
    lr.font.color.rgb = GREY
    lr.font.name  = "Calibri"

    title_p = doc.add_heading(req.get("component_name", "Engineering Component"), 0)
    title_r = title_p.runs[0] if title_p.runs else title_p.add_run()
    title_r.font.color.rgb = DARK_BLUE
    title_r.font.name = "Calibri Bold"

    sub_p = doc.add_paragraph()
    sub_r = sub_p.add_run(
        f"Report ID: {rep.get('report_id', 'N/A')}     |     "
        f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}"
    )
    sub_r.font.size  = Pt(10)
    sub_r.font.color.rgb = BLUE
    sub_r.font.name  = "Calibri"
    doc.add_paragraph()

    # Executive Summary
    exec_summary = rep.get("executive_summary", "")
    if exec_summary:
        add_heading(doc, "Executive Summary", 1)
        add_body(doc, exec_summary)
        doc.add_paragraph()

    # ── Section 1: Requirements ────────────────────────────────────────
    add_heading(doc, "1. Engineering Requirements", 1)

    req_data = [
        ["Field", "Value"],
        ["Component", req.get("component_name", "")],
        ["Priority", req.get("priority", "").upper()],
        ["Design Space", req.get("design_space", "Not specified")],
    ]
    add_two_col_table(doc, req_data)
    doc.add_paragraph()

    func_reqs = req.get("functional_requirements", [])
    if func_reqs:
        add_heading(doc, "Functional Requirements", 2)
        for r in func_reqs:
            add_bullet(doc, str(r))

    perf = req.get("performance_targets", {})
    if perf:
        add_heading(doc, "Performance Targets", 2)
        for k, v in perf.items():
            add_bullet(doc, f"{k}: {v}")

    constraints = req.get("constraints", [])
    if constraints:
        add_heading(doc, "Constraints", 2)
        for c in constraints:
            add_bullet(doc, str(c))

    # ── Section 2: Design Alternatives ────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "2. Design Alternatives", 1)

    rec_id = ev.get("recommended", "B")

    for alt in alts:
        is_rec    = alt["id"] == rec_id
        alt_title = f"Alternative {alt['id']}: {alt.get('name', '')}"
        if is_rec:
            alt_title += "   ★ RECOMMENDED"

        add_heading(doc, alt_title, 2)

        score = ev.get("scores", {}).get(alt["id"], "N/A")
        risk  = ev.get("risk_assessment", {}).get(alt["id"], "MEDIUM")

        alt_data = [
            ["Metric", "Value"],
            ["Material",             alt.get("material", "")],
            ["Manufacturing Process", alt.get("manufacturing_process", "")],
            ["Estimated Weight",     f"{alt.get('estimated_weight_kg', '')} kg"],
            ["Cost Index",           f"{alt.get('estimated_cost_index', '')}/5"],
            ["Performance Score",    f"{alt.get('performance_score', '')}/5"],
            ["Sustainability Score", f"{alt.get('sustainability_score', '')}/5"],
            ["Feasibility",          alt.get("feasibility", "")],
            ["Overall Score",        f"{score}/100"],
            ["Risk",                 risk],
        ]
        bg = "D5E8F0" if is_rec else "F0F4F8"
        add_two_col_table(doc, alt_data, row_bgs=(bg, "FFFFFF"))
        doc.add_paragraph()

        add_body(doc, alt.get("concept", ""))

        advantages = alt.get("advantages", [])
        if advantages:
            p = doc.add_paragraph()
            p.add_run("Advantages: ").font.bold = True
            p.runs[0].font.name = "Calibri"
            p.runs[0].font.size = Pt(10)
            p.add_run(clean(" · ".join(str(a) for a in advantages))).font.name = "Calibri"

        disadvantages = alt.get("disadvantages", [])
        if disadvantages:
            p = doc.add_paragraph()
            p.add_run("Disadvantages: ").font.bold = True
            p.runs[0].font.name = "Calibri"
            p.runs[0].font.size = Pt(10)
            p.add_run(clean(" · ".join(str(d) for d in disadvantages))).font.name = "Calibri"

        doc.add_paragraph()

    # ── Section 3: Evaluation & Recommendation ────────────────────────
    doc.add_page_break()
    add_heading(doc, "3. Evaluation & Recommendation", 1)

    ranking = ev.get("ranking", [])
    if ranking:
        add_body(doc, f"Ranking: {' > '.join(ranking)}")

    rationale = ev.get("recommendation_rationale", "")
    if rationale:
        add_body(doc, f"Recommendation: Alternative {rec_id}", bold=True, color=DARK_BLUE)
        add_body(doc, rationale)
        doc.add_paragraph()

    next_steps = rep.get("next_steps", [])
    if next_steps:
        add_heading(doc, "Next Steps", 2)
        for i, step in enumerate(next_steps, 1):
            add_bullet(doc, f"{i}. {clean(str(step))}")

    sources = rep.get("knowledge_sources", [])
    if sources:
        doc.add_paragraph()
        add_heading(doc, "Knowledge Sources", 2)
        for s in sources:
            add_bullet(doc, str(s))

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    fr = footer_p.add_run(
        f"Generated by Generative Design Assistant v1.0  |  "
        f"{datetime.now().strftime('%d %B %Y')}"
    )
    fr.font.size  = Pt(8.5)
    fr.font.color.rgb = GREY
    fr.font.name  = "Calibri"

    doc.save(output_path)
    logger.success(f"Word document exported: {output_path}")
    return output_path